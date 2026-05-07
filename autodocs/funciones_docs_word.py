import io
import os
import re
import tempfile
import traceback
import unicodedata
from collections import Counter

import pandas as pd
import openpyxl

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from google.oauth2 import service_account
from google.oauth2 import credentials as oauth2_credentials
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from googleapiclient.errors import HttpError


# ==============================
# TAMAÑO POR DEFECTO DE IMÁGENES
# ==============================
DEFAULT_IMAGE_WIDTH_PT  = 300
DEFAULT_IMAGE_HEIGHT_PT = 200

# ==============================
# MAPEO DE TIPOS DESDE EL EXCEL
# ==============================
TIPO_MAP = {
    "texto":          "text",
    "número":         "text",
    "numero":         "text",
    "condicional":    "text",
    "imagen":         "image",
    "tabla":          "table",
    "loop":           "loop",
    "imagen / tabla": "image",
}

# ==============================
# VERBOSE / LOGGING
# ==============================
VERBOSE = False

def _log(*args, **kwargs):
    if VERBOSE:
        print(*args, **kwargs)

def _debug(msg):
    log_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "debug_tablas.log")
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(msg + "\n")


# ==============================
# GOOGLE AUTH
# ==============================

SCOPES = [
    'https://www.googleapis.com/auth/drive',
    'https://www.googleapis.com/auth/spreadsheets',
]


def autenticar_oauth(ruta_client_secret=None, ruta_token=None):
    """
    Autenticación OAuth2. Devuelve (None, drive_service, sheets_service).
    docs_service se retorna como None por compatibilidad con doc_filler.py.
    """
    if ruta_client_secret is None:
        ruta_client_secret = os.getenv("GOOGLE_OAUTH_CLIENT")
        if not ruta_client_secret:
            raise Exception("Define GOOGLE_OAUTH_CLIENT con la ruta al client_secret.json")

    if not os.path.exists(ruta_client_secret):
        raise FileNotFoundError(f"No existe: {ruta_client_secret}")

    if ruta_token is None:
        ruta_token = os.getenv(
            "GOOGLE_OAUTH_TOKEN",
            os.path.join(os.path.dirname(ruta_client_secret), "token.json")
        )

    creds = None
    if os.path.exists(ruta_token):
        creds = oauth2_credentials.Credentials.from_authorized_user_file(ruta_token, SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(ruta_client_secret, SCOPES)
            creds = flow.run_local_server(port=0)
        with open(ruta_token, "w") as f:
            f.write(creds.to_json())

    drive_service  = build('drive',  'v3', credentials=creds)
    sheets_service = build('sheets', 'v4', credentials=creds)
    return None, drive_service, sheets_service


def autenticar_servicio(ruta_credenciales):
    """Autenticación con service account."""
    creds = service_account.Credentials.from_service_account_file(ruta_credenciales, scopes=SCOPES)
    drive_service  = build('drive',  'v3', credentials=creds)
    sheets_service = build('sheets', 'v4', credentials=creds)
    return None, drive_service, sheets_service


# ==============================
# OPERACIONES DE DRIVE
# ==============================

def extraer_id_gdoc(url_o_id):
    match = re.search(r"/(?:d|folders)/([a-zA-Z0-9_-]+)", url_o_id)
    if match:
        return match.group(1)
    return url_o_id.strip()


def descargar_excel_drive(file_id, drive_service):
    """Descarga un Excel de Drive y lo retorna como BytesIO."""
    request = drive_service.files().get_media(fileId=file_id)
    fh = io.BytesIO()
    dl = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        _, done = dl.next_chunk()
    fh.seek(0)
    return fh


def descargar_docx_drive(file_id, drive_service):
    """
    Descarga un .docx de Drive a un temporal local y retorna la ruta.
    Si es un Google Doc nativo, lo exporta como .docx automáticamente.
    """
    meta = drive_service.files().get(fileId=file_id, fields="mimeType,name").execute()
    mime = meta.get("mimeType", "")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()

    if mime == "application/vnd.google-apps.document":
        _log(f"   Exportando Google Doc '{meta['name']}' como .docx...")
        response = drive_service.files().export_media(
            fileId=file_id,
            mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        _log(f"   Descargando '{meta['name']}'...")
        response = drive_service.files().get_media(fileId=file_id)

    with open(tmp.name, "wb") as f:
        dl = MediaIoBaseDownload(f, response)
        done = False
        while not done:
            _, done = dl.next_chunk()

    _log(f"   .docx descargado -> {tmp.name}")
    return tmp.name


def subir_docx_drive(ruta_local, nombre, carpeta_id, drive_service):
    """Sube el .docx local a Drive y retorna el ID del archivo creado."""
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    file_meta = {"name": nombre, "parents": [carpeta_id]}
    media = MediaFileUpload(ruta_local, mimetype=mime, resumable=True)
    result = drive_service.files().create(body=file_meta, media_body=media, fields="id").execute()
    file_id = result.get("id")
    _log(f"   Archivo subido a Drive -> ID: {file_id}")
    return file_id


def copiar_documento(doc_id, nombre_nuevo, drive_service, carpeta_destino_id=None):
    """
    Descarga la plantilla .docx a un temporal local.
    Retorna la RUTA LOCAL del archivo (no copia en Drive).
    carpeta_destino_id se ignora aqui; se usa al subir el resultado final.
    """
    return descargar_docx_drive(doc_id, drive_service)


def listar_archivos_drive(drive_service, page_size=10):
    results = drive_service.files().list(pageSize=page_size, fields="files(id, name)").execute()
    return results.get('files', [])


# ==============================
# CARGA DEL DICCIONARIO
# ==============================

def cargar_diccionario(archivo, project_filter=None):
    """
    Lee el Excel del diccionario. Logica identica al original.
    Acepta ruta (str) o BytesIO.
    """
    if isinstance(archivo, (str, bytes, os.PathLike)):
        df = pd.read_csv(archivo) if str(archivo).endswith(".csv") else pd.read_excel(archivo, header=None)
    else:
        df = pd.read_excel(archivo, header=None)

    header_idx = 0
    for i, row in df.iterrows():
        if any(str(v).strip().lower() == "nombre del dato" for v in row.values):
            header_idx = i
            break

    df.columns = df.iloc[header_idx]
    df = df[header_idx + 1:].reset_index(drop=True)
    df.columns = [str(col).strip().lower() for col in df.columns]

    conteo = Counter(df.columns)
    duplicadas = [col for col, n in conteo.items() if n > 1]
    if duplicadas:
        print(f"Columnas duplicadas: {duplicadas}")
    else:
        print("No hay columnas duplicadas")
    print(f"   Columnas: {list(df.columns)}")

    df = df.rename(columns={
        "nombre del dato": "placeholder",
        "valor":           "value",
        "tipo":            "type",
        "alias":           "alias",
    })

    tiene_width      = "width_pt"   in df.columns
    tiene_height     = "height_pt"  in df.columns
    tiene_sheet      = "sheet"      in df.columns
    tiene_header_row = "header_row" in df.columns

    data = {}

    for _, row in df.iterrows():
        placeholder = str(row.get("placeholder", "")).strip()
        raw_value   = row.get("value")
        raw_type    = str(row.get("type", "")).strip().lower()
        alias       = str(row.get("alias", "")).strip()

        if not placeholder and not alias:
            continue

        key_limpia = limpiar_placeholder(placeholder)
        final_key  = alias if alias else key_limpia
        tipo_norm  = TIPO_MAP.get(raw_type, "text")

        if tipo_norm == "text":
            value = _convertir_a_texto(raw_value)
        elif tipo_norm == "image":
            file_id   = str(raw_value).strip() if not _es_nulo(raw_value) else None
            width_pt  = _leer_dim(row, "width_pt")  if tiene_width  else DEFAULT_IMAGE_WIDTH_PT
            height_pt = _leer_dim(row, "height_pt") if tiene_height else DEFAULT_IMAGE_HEIGHT_PT
            value = {"file_id": file_id, "width_pt": width_pt, "height_pt": height_pt}
        elif tipo_norm == "table":
            raw_url    = str(raw_value).strip() if not _es_nulo(raw_value) else None
            sheet      = str(row.get("sheet", "")).strip() if tiene_sheet else None
            sheet      = sheet if sheet and sheet.lower() != "nan" else None
            header_row = _leer_header_row(row, tiene_header_row)
            value = {"file_id": raw_url, "sheet": sheet, "header_row": header_row} if raw_url else None
        elif tipo_norm == "loop":
            raw_url    = str(raw_value).strip() if not _es_nulo(raw_value) else None
            prefijo    = str(row.get("sheet", "")).strip() if tiene_sheet else None
            prefijo    = prefijo if prefijo and prefijo.lower() != "nan" else None
            header_row = _leer_header_row(row, tiene_header_row)
            value = {"file_id": raw_url, "prefijo": prefijo, "header_row": header_row} if raw_url else None
        else:
            value = _convertir_a_texto(raw_value)

        data[final_key] = {"placeholder": placeholder, "type": tipo_norm, "value": value}

    return data


# ==============================
# HELPERS GENERALES
# ==============================

def limpiar_placeholder(texto):
    return re.sub(r"[{}]", "", texto).strip()

def _es_nulo(valor):
    try:
        return pd.isna(valor)
    except Exception:
        return valor is None or str(valor).strip() == ""

def _convertir_a_texto(valor):
    if _es_nulo(valor):
        return None
    if isinstance(valor, float) and valor == int(valor):
        return str(int(valor))
    return str(valor)

def _leer_dim(row, col_name):
    val = row.get(col_name)
    if _es_nulo(val):
        return DEFAULT_IMAGE_WIDTH_PT if col_name == "width_pt" else DEFAULT_IMAGE_HEIGHT_PT
    try:
        return int(float(val))
    except Exception:
        return DEFAULT_IMAGE_WIDTH_PT if col_name == "width_pt" else DEFAULT_IMAGE_HEIGHT_PT

def _leer_header_row(row, tiene_col):
    if not tiene_col:
        return None
    hr = row.get("header_row")
    if _es_nulo(hr):
        return None
    try:
        return int(float(hr))
    except Exception:
        return None

def _evaluar_booleano(valor):
    if valor is None:
        return False
    if isinstance(valor, bool):
        return valor
    return str(valor).strip().lower() in ("true", "1", "si", "yes", "si")

def _hex_argb_a_rgb(hex_argb):
    h = hex_argb[-6:]
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


# ==============================
# HELPERS python-docx
# ==============================

def _iter_all_paragraphs(doc):
    """Itera todos los parrafos del documento incluyendo dentro de tablas."""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para

def _get_para_full_text(para):
    return "".join(r.text for r in para.runs)

def _consolidar_runs(para):
    runs = para.runs
    if len(runs) <= 1:
        return
    texto_completo = "".join(r.text for r in runs)
    if not texto_completo:
        return
    runs[0].text = texto_completo
    for r in runs[1:]:
        r.text = ""

def _replace_text_in_para(para, old, new):
    _consolidar_runs(para)
    full = _get_para_full_text(para)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_full)
    return True


# ==============================
# CONDICIONALES  {% if %}...{% endif %}
# ==============================

def procesar_condicionales(doc_path, diccionario, _docs_service=None):
    """
    Evalua bloques {% if variable %}...{% endif %} en el .docx local.
    Soporta {% else %} opcional.
    - Si la variable es True  -> elimina solo las etiquetas, conserva contenido
    - Si la variable es False -> elimina el bloque completo
    """
    doc  = Document(doc_path)
    body = doc.element.body

    def _all_body_paras(parent):
        result = []
        for child in parent:
            local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if local == "p":
                result.append(child)
            else:
                result.extend(_all_body_paras(child))
        return result

    def _para_text(el):
        return "".join((t.text or "") for t in el.iter(qn("w:t")))

    procesados = []
    omitidos   = []

    for _ in range(200):  # max iteraciones de seguridad
        paras  = _all_body_paras(body)
        idx_if = None
        var    = None

        for i, p in enumerate(paras):
            m = re.search(r"\{%-?\s*if\s+(\w+)\s*-?%\}", _para_text(p))
            if m:
                idx_if = i
                var    = m.group(1).strip()
                break

        if idx_if is None:
            break

        entrada = diccionario.get(var)
        if entrada is None:
            omitidos.append({"variable": var, "razon": "no encontrada en el diccionario"})
            paras[idx_if].getparent().remove(paras[idx_if])
            continue

        es_verdadero = _evaluar_booleano(entrada.get("value"))

        idx_else  = None
        idx_endif = None
        for j in range(idx_if + 1, len(paras)):
            t = _para_text(paras[j])
            if re.search(r"\{%-?\s*else\s*-?%\}", t) and idx_else is None:
                idx_else = j
            if re.search(r"\{%-?\s*endif\s*-?%\}", t):
                idx_endif = j
                break

        if idx_endif is None:
            paras[idx_if].getparent().remove(paras[idx_if])
            continue

        paras_if_content   = paras[idx_if + 1 : idx_else if idx_else is not None else idx_endif]
        paras_else_content = paras[idx_else + 1 : idx_endif] if idx_else is not None else []

        if es_verdadero:
            paras[idx_if].getparent().remove(paras[idx_if])
            if idx_else is not None:
                paras[idx_else].getparent().remove(paras[idx_else])
                for p in paras_else_content:
                    p.getparent().remove(p)
            paras[idx_endif].getparent().remove(paras[idx_endif])
        else:
            paras[idx_if].getparent().remove(paras[idx_if])
            for p in paras_if_content:
                p.getparent().remove(p)
            if idx_else is not None:
                paras[idx_else].getparent().remove(paras[idx_else])
            paras[idx_endif].getparent().remove(paras[idx_endif])

        procesados.append({
            "variable":     var,
            "valor":        entrada.get("value"),
            "es_verdadero": es_verdadero,
            "tiene_else":   idx_else is not None,
        })

    doc.save(doc_path)
    _log(f"Condicionales procesadas: {len(procesados)}")
    for p in procesados:
        icono = "V" if p["es_verdadero"] else "X"
        else_str = " (con else)" if p["tiene_else"] else ""
        _log(f"   {icono} if {p['variable']} -> {p['valor']}{else_str}")
    return {"procesados": procesados, "omitidos": omitidos}


# ==============================
# REEMPLAZO DE TEXTO
# ==============================

def reemplazar_textos(doc_path, diccionario, _docs_service=None):
    """Reemplaza todos los placeholders de tipo text en el .docx local."""
    doc          = Document(doc_path)
    reemplazados = []
    omitidos     = []

    entradas = [(a, e) for a, e in diccionario.items() if e.get("type") == "text"]

    if not entradas:
        _log("No se encontraron entradas de tipo text.")
        return {"reemplazados": [], "omitidos": omitidos}

    for alias, entrada in entradas:
        placeholder     = entrada.get("placeholder", "").strip()
        value           = entrada.get("value")
        texto_reemplazo = "" if value is None else str(value)

        if not placeholder:
            omitidos.append({"alias": alias, "razon": "placeholder vacio"})
            continue

        ocurrencias = sum(
            1 for para in _iter_all_paragraphs(doc)
            if _replace_text_in_para(para, placeholder, texto_reemplazo)
        )

        if ocurrencias > 0:
            reemplazados.append({
                "alias": alias, "placeholder": placeholder,
                "value": texto_reemplazo, "ocurrencias": ocurrencias,
            })
            _log(f"   {placeholder} -> '{texto_reemplazo}' ({ocurrencias} vez/veces)")
        else:
            omitidos.append({"alias": alias, "razon": "placeholder no encontrado en el documento"})

    doc.save(doc_path)
    _log(f"Texto reemplazado: {len(reemplazados)} placeholders")
    return {"reemplazados": reemplazados, "omitidos": omitidos}


# ==============================
# REEMPLAZO DE IMAGENES
# ==============================

def _descargar_imagen_drive(file_id, drive_service):
    meta = drive_service.files().get(fileId=file_id, fields="mimeType").execute()
    ext  = {"image/png": ".png", "image/jpeg": ".jpg",
            "image/gif": ".gif", "image/webp": ".webp", "image/bmp": ".bmp"}.get(
        meta.get("mimeType", ""), ".png")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    tmp.close()

    def _dl(request):
        fh = io.BytesIO()
        dl = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            _, done = dl.next_chunk()
        return fh.getvalue()

    try:
        data = _dl(drive_service.files().get_media(fileId=file_id))
    except HttpError:
        data = _dl(drive_service.files().export_media(fileId=file_id, mimeType="image/png"))

    with open(tmp.name, "wb") as f:
        f.write(data)
    return tmp.name


def reemplazar_imagenes(doc_path, diccionario, _docs_service=None, drive_service=None):
    """
    Reemplaza placeholders de tipo image en el .docx local.
    Descarga la imagen de Drive e inserta un InlineImage en el parrafo.
    Respeta width_pt y height_pt del diccionario.
    """
    doc          = Document(doc_path)
    reemplazados = []
    omitidos     = []

    entradas = [(a, e) for a, e in diccionario.items() if e.get("type") == "image"]

    if not entradas:
        _log("No se encontraron entradas de tipo image.")
        return {"reemplazados": [], "omitidos": omitidos}

    for alias, entrada in entradas:
        placeholder = entrada.get("placeholder", "").strip()
        valor       = entrada.get("value", {})
        raw_file_id = valor.get("file_id")
        width_pt    = valor.get("width_pt",  DEFAULT_IMAGE_WIDTH_PT)
        height_pt   = valor.get("height_pt", DEFAULT_IMAGE_HEIGHT_PT)

        if not raw_file_id:
            omitidos.append({"alias": alias, "razon": "file_id vacio o no definido"})
            continue
        if not placeholder:
            omitidos.append({"alias": alias, "razon": "placeholder vacio"})
            continue

        file_id = extraer_id_gdoc(raw_file_id)

        try:
            img_path = _descargar_imagen_drive(file_id, drive_service)
        except Exception as e:
            omitidos.append({"alias": alias, "razon": f"error descargando imagen: {e}"})
            continue

        encontrado = False
        for para in _iter_all_paragraphs(doc):
            _consolidar_runs(para)
            if placeholder not in _get_para_full_text(para):
                continue
            encontrado = True
            for run in para.runs:
                run.text = ""
            run = para.runs[0] if para.runs else para.add_run()
            run.add_picture(img_path, width=Pt(width_pt), height=Pt(height_pt))

        if encontrado:
            reemplazados.append({
                "alias": alias, "placeholder": placeholder,
                "file_id": file_id, "width_pt": width_pt, "height_pt": height_pt,
            })
            _log(f"   {placeholder} -> imagen {width_pt}x{height_pt} pt")
        else:
            omitidos.append({"alias": alias, "razon": "placeholder no encontrado en el documento"})

    doc.save(doc_path)
    _log(f"Imagenes reemplazadas: {len(reemplazados)}")
    omitidos_reales = [o for o in omitidos if "no encontrado" not in o.get("razon", "")]
    if omitidos_reales:
        _log(f"Imagenes omitidas: {len(omitidos_reales)}")
        for o in omitidos_reales:
            _log(f"   {o['alias']}: {o['razon']}")
    return {"reemplazados": reemplazados, "omitidos": omitidos}


# ==============================
# HELPERS CARGA DE EXCEL
# ==============================

def _cargar_datos_tabla(url, sheet_name, drive_service, header_row=None, alias=None, placeholder=None):
    """
    Descarga un .xlsx de Drive y extrae datos, formato de celda y merges.
    Logica identica a la version original.
    Retorna (df, cell_formats, merges).
    """
    file_id = extraer_id_gdoc(url)

    def _descargar(request):
        fh = io.BytesIO()
        dl = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            _, done = dl.next_chunk()
        fh.seek(0)
        return fh

    try:
        fh = _descargar(drive_service.files().get_media(fileId=file_id))
    except HttpError as e:
        if e.resp.status in (400, 403):
            fh = _descargar(drive_service.files().export_media(
                fileId=file_id,
                mimeType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ))
        else:
            raise

    fh.seek(0)
    wb = openpyxl.load_workbook(fh, data_only=True)
    n_hojas = len(wb.sheetnames)

    if sheet_name and sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        hoja_usada = sheet_name
    elif n_hojas == 1:
        ws = wb.active
        hoja_usada = ws.title
    else:
        candidato = _resolver_hoja_por_nombre(wb.sheetnames, placeholder, alias)
        if candidato:
            ws = wb[candidato]
            hoja_usada = candidato
            _log(f"   No se especifico sheet para '{placeholder or alias}'. Usando '{candidato}'.")
        else:
            hojas_disp = ", ".join(f"'{h}'" for h in wb.sheetnames)
            etiqueta   = placeholder or alias or "(sin etiqueta)"
            raise ValueError(
                f"El Excel para '{etiqueta}' tiene {n_hojas} hojas pero no se indico cual usar. "
                f"Hojas: {hojas_disp}. Agrega 'sheet' en el diccionario."
            )

    _log(f"   Hoja: '{hoja_usada}' ({ws.max_row} x {ws.max_column})")

    max_col_real = max_row_real = 0
    first_row_con_datos = None
    for fila in ws.iter_rows():
        for celda in fila:
            val = celda.value
            if val is not None and str(val).strip() not in ("", "'"):
                max_col_real = max(max_col_real, celda.column)
                max_row_real = max(max_row_real, celda.row)
                if first_row_con_datos is None:
                    first_row_con_datos = celda.row
    if max_col_real == 0:
        max_col_real = ws.max_column
    if max_row_real == 0:
        max_row_real = ws.max_row

    if header_row and header_row > 0:
        header_idx = header_row - 1
        fila_c = list(ws.iter_rows(min_row=header_idx+1, max_row=header_idx+1, max_col=max_col_real, values_only=False))
        if fila_c:
            tiene_cont = any(c.value is not None and str(c.value).strip() not in ("", "'") for c in fila_c[0])
            if not tiene_cont and first_row_con_datos is not None:
                header_idx = first_row_con_datos - 1
    else:
        header_idx = (first_row_con_datos - 1) if first_row_con_datos else 0

    filas_excel = list(ws.iter_rows(
        min_row=header_idx + 1, max_row=max_row_real,
        max_col=max_col_real, values_only=False
    ))

    def _limpiar_enc(val):
        if val is None:
            return ""
        s = str(val)
        return s.lstrip("'") if s.startswith("'") else s

    encabezados = [_limpiar_enc(c.value) for c in filas_excel[0]]
    datos_filas = [
        [str(c.value) if c.value is not None else "" for c in fila]
        for fila in filas_excel[1:]
    ]
    df = pd.DataFrame(datos_filas, columns=encabezados).fillna("").astype(str)

    # Colores ARGB que consideramos "sin relleno real" (transparente / blanco / negro)
    _FILLS_IGNORADOS = {"00000000", "FFFFFFFF", "FF000000", "00FFFFFF", "FFFFFFFF"}

    cell_formats = {}
    for doc_ri, fila in enumerate(filas_excel):
        for ci, celda in enumerate(fila):
            if ci >= max_col_real:
                break

            has_value = (
                celda.value is not None
                and str(celda.value).strip() not in ("", "'")
            )

            fmt = {}

            # ── Fondo: se captura para TODAS las celdas (incluso vacías).
            # Esto preserva el relleno azul de celdas de cabecera que no tienen
            # texto pero sí color (p.ej. A2:D2 en EOLOVANOS).
            fill = celda.fill
            if fill and fill.fill_type not in (None, "none"):
                color = fill.fgColor
                if color and color.type == "rgb" and color.rgb not in _FILLS_IGNORADOS:
                    fmt["bg_color"] = color.rgb

            # ── Fuente: solo para celdas con contenido real
            if has_value:
                if celda.font and celda.font.bold:
                    fmt["bold"] = True
                if celda.font and celda.font.size:
                    fmt["font_size"] = float(celda.font.size)
                    _debug(
                        f"_cargar_datos_tabla | {celda.coordinate} "
                        f"| valor={celda.value!r} | font_size={fmt['font_size']}"
                    )
                if celda.font and celda.font.color:
                    fc = celda.font.color
                    if fc.type == "rgb" and fc.rgb not in ("00000000", "FF000000"):
                        fmt["font_color"] = fc.rgb

            if fmt:
                cell_formats[(doc_ri, ci)] = fmt

    merges = []
    for rng in ws.merged_cells.ranges:
        doc_min_row = rng.min_row - 1 - header_idx
        doc_max_row = rng.max_row - 1 - header_idx
        min_col     = rng.min_col - 1
        max_col     = min(rng.max_col - 1, max_col_real - 1)
        if doc_min_row >= 0 and min_col < max_col_real:
            merges.append({
                "min_row":  doc_min_row, "max_row":  doc_max_row,
                "min_col":  min_col,     "max_col":  max_col,
                "row_span": doc_max_row - doc_min_row + 1,
                "col_span": max_col - min_col + 1,
            })

    return df, cell_formats, merges


def _resolver_hoja_por_nombre(hojas_disponibles, placeholder, alias):
    STOPWORDS = {"de","del","la","el","los","las","y","en","con","sobre","por","para","a","al","un","una","mt","ac","dc"}

    def _normalizar(s):
        if not s:
            return ""
        s = str(s)
        s = "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")
        return re.sub(r"[^a-zA-Z0-9]+", " ", s).lower().strip()

    def _tokens(s):
        return {t for t in _normalizar(s).split() if t and t not in STOPWORDS}

    objetivo = _tokens(placeholder) | _tokens(alias)
    if not objetivo:
        return None

    mejor_hoja, mejor_score = None, 0
    for hoja in hojas_disponibles:
        toks = _tokens(hoja)
        if not toks:
            continue
        score = len(objetivo & toks) / max(len(toks), 1)
        if score > mejor_score:
            mejor_score = score
            mejor_hoja  = hoja

    return mejor_hoja if mejor_score >= 0.5 else None


# ==============================
# CREAR TABLA DOCX CON FORMATO
# ==============================

def _aplicar_bordes_tabla(tabla):
    """Aplica bordes de cuadrícula via XML, sin depender del estilo 'Table Grid'."""
    tbl   = tabla._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old_b in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old_b)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _quitar_bordes_celda(tc):
    """Quita todos los bordes de una celda individual."""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    for old_b in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old_b)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _vaciar_celda_xml(tc):
    """
    Elimina todo el contenido de una celda dejando exactamente UN párrafo vacío.
    Word requiere al menos un <w:p> por celda. Múltiples <w:p> vacíos producen
    saltos de línea que inflan el alto de fila.
    """
    parrafos = tc.findall(qn("w:p"))
    for p in parrafos:
        for r in list(p.findall(qn("w:r"))):
            p.remove(r)
        for hyp in list(p.findall(qn("w:hyperlink"))):
            p.remove(hyp)
    for p in parrafos[1:]:
        tc.remove(p)
    if not parrafos:
        tc.append(OxmlElement("w:p"))


def _aplicar_formato_tabla(tabla, cell_formats, merges):
    """
    Aplica merges y formato de celda (bg, bold, font_size, font_color).

    IMPORTANTE – comportamiento de python-docx al fusionar celdas:
    Cuando se llama a cell.merge(), los párrafos de cada celda absorbida
    se MUEVEN al interior del anchor (celda origen del merge), lo que genera
    párrafos vacíos extra dentro de esa celda y produce grandes espacios en
    blanco. La función _limpiar_parrafos_merge() corrige esto después.
    """
    # ── 1. Aplicar merges
    for m in sorted(merges, key=lambda x: (-x["min_row"], -x["min_col"])):
        try:
            tabla.cell(m["min_row"], m["min_col"]).merge(
                tabla.cell(m["max_row"], m["max_col"])
            )
        except Exception:
            pass

    # ── 2. Aplicar formato celda a celda
    for (fi, ci), fmt in cell_formats.items():
        try:
            celda = tabla.cell(fi, ci)
        except Exception:
            continue

        # Fondo: se aplica incluso a celdas sin texto (p.ej. A2:D2 azul en EOLOVANOS)
        if "bg_color" in fmt:
            tc_pr = celda._tc.get_or_add_tcPr()
            # Eliminar shd previos para evitar duplicados
            for old_shd in tc_pr.findall(qn("w:shd")):
                tc_pr.remove(old_shd)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  fmt["bg_color"][-6:])
            tc_pr.append(shd)

        # Fuente: solo si la celda tiene texto real
        texto_celda = "".join(
            t.text or "" for t in celda._tc.iter(qn("w:t"))
        ).strip()
        if texto_celda:
            for para in celda.paragraphs:
                for run in para.runs:
                    if fmt.get("bold"):
                        run.bold = True
                    if "font_size" in fmt:
                        run.font.size = Pt(fmt["font_size"])
                    if "font_color" in fmt:
                        r, g, b = _hex_argb_a_rgb(fmt["font_color"])
                        run.font.color.rgb = RGBColor(r, g, b)


def _limpiar_parrafos_merge(tabla):
    """
    Elimina los párrafos vacíos extra que python-docx introduce en la celda
    anchor cuando fusiona celdas con merge().

    Problema confirmado: al llamar a cell.merge(), python-docx MUEVE los
    párrafos de cada celda absorbida al interior del anchor. Si esas celdas
    tenían texto vacío (""), el anchor acaba con N párrafos vacíos de más
    → espacio en blanco gigante dentro de la celda combinada.

    Solución: después del merge, recorrer todos los <w:tc> y eliminar los
    párrafos vacíos que siguen al último párrafo con contenido real,
    dejando siempre al menos un <w:p> (requisito de Word).
    """
    for row in tabla.rows:
        for tc in row._tr.findall(qn("w:tc")):
            parrafos = tc.findall(qn("w:p"))
            if len(parrafos) <= 1:
                continue  # nada que limpiar

            # Eliminar párrafos vacíos desde el final hacia atrás
            while len(parrafos) > 1:
                ultimo = parrafos[-1]
                texto = "".join(
                    (t.text or "") for t in ultimo.iter(qn("w:t"))
                ).strip()
                if texto:
                    break          # encontramos contenido real → parar
                tc.remove(ultimo)
                parrafos.pop()


def _hacer_filas_asimetricas(tabla):
    """
    Crea una tabla verdaderamente asimétrica en OOXML eliminando las celdas
    vacías del FINAL de cada fila y ampliando la última celda con datos
    mediante gridSpan.

    Por qué val="none" en tcBorders NO es suficiente:
    En el modelo de resolución de conflictos de OOXML, cuando una celda vacía
    tiene val="none" en su borde izquierdo pero la celda adyacente con DATOS
    no tiene override explícito (usa el borde de tabla insideV="single"), el
    borde "single" de la celda con datos gana y se dibuja igual.
    La única forma de eliminar el borde completamente es quitar la celda vacía
    del DOM y ampliar la última celda con datos.

    Qué hace esta función:
    - Para cada fila, localiza la última celda con texto real.
    - Calcula cuántas celdas vacías hay a continuación (trailing).
    - Suma su gridSpan a la última celda con datos y actualiza su ancho (tcW).
    - Elimina las celdas vacías trailing del <w:tr>.

    Celdas vacías NO-trailing (antes de la última con datos) se dejan tal cual;
    son manejadas por _ocultar_celdas_vacias_intermedias().
    """
    tbl = tabla._tbl

    # Leer anchos de columna del tblGrid
    tblGrid = tbl.find(qn("w:tblGrid"))
    grid_widths: list[int] = []
    if tblGrid is not None:
        for gridCol in tblGrid.findall(qn("w:gridCol")):
            w = gridCol.get(qn("w:w"))
            grid_widths.append(int(w) if w else 0)

    for row in tabla.rows:
        tcs = list(row._tr.findall(qn("w:tc")))
        if not tcs:
            continue

        # Mapear cada <w:tc> a (índice_físico, col_lógica_inicio, gridSpan)
        tc_info: list[tuple] = []
        col_logica = 0
        for tc in tcs:
            tcPr = tc.find(qn("w:tcPr"))
            gs = 1
            if tcPr is not None:
                gs_el = tcPr.find(qn("w:gridSpan"))
                if gs_el is not None:
                    gs = int(gs_el.get(qn("w:val"), 1))
            tc_info.append((tc, col_logica, gs))
            col_logica += gs

        # Hallar el último índice físico que tiene texto
        ultimo_con_datos = -1
        for i, (tc, _, _) in enumerate(tc_info):
            texto = "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip()
            if texto:
                ultimo_con_datos = i

        # Sin datos o sin trailing vacíos → nada que hacer
        if ultimo_con_datos < 0 or ultimo_con_datos >= len(tc_info) - 1:
            continue

        celdas_trailing = tc_info[ultimo_con_datos + 1:]
        gs_extra = sum(gs for _, _, gs in celdas_trailing)
        if gs_extra == 0:
            continue

        # ── Ampliar la última celda con datos ──
        ultima_tc, ultima_col, ultima_gs = tc_info[ultimo_con_datos]
        nueva_gs = ultima_gs + gs_extra

        ultima_tcPr = ultima_tc.find(qn("w:tcPr"))
        if ultima_tcPr is None:
            ultima_tcPr = OxmlElement("w:tcPr")
            ultima_tc.insert(0, ultima_tcPr)

        # Actualizar o crear gridSpan
        gs_el_existente = ultima_tcPr.find(qn("w:gridSpan"))
        if gs_el_existente is not None:
            ultima_tcPr.remove(gs_el_existente)
        if nueva_gs > 1:
            new_gs_el = OxmlElement("w:gridSpan")
            new_gs_el.set(qn("w:val"), str(nueva_gs))
            tcW_el = ultima_tcPr.find(qn("w:tcW"))
            if tcW_el is not None:
                tcW_el.addnext(new_gs_el)
            else:
                ultima_tcPr.insert(0, new_gs_el)

        # Actualizar ancho de celda para que sume los anchos de las columnas absorbidas
        if grid_widths:
            fin_col = ultima_col + nueva_gs
            new_width = sum(
                grid_widths[c]
                for c in range(ultima_col, min(fin_col, len(grid_widths)))
            )
            tcW_el = ultima_tcPr.find(qn("w:tcW"))
            if tcW_el is None:
                tcW_el = OxmlElement("w:tcW")
                ultima_tcPr.insert(0, tcW_el)
            tcW_el.set(qn("w:type"), "dxa")
            tcW_el.set(qn("w:w"), str(new_width))

        # ── Eliminar las celdas vacías trailing del <w:tr> ──
        for tc, _, _ in celdas_trailing:
            row._tr.remove(tc)


def _ocultar_celdas_vacias_intermedias(tabla):
    """
    Para las celdas vacías que NO están al final de la fila (no-trailing)
    — como las celdas A2:D2 vacías en la cabecera de EOLOVANOS, que están
    ANTES de la celda con datos "Gravivano (m)" — aplica val="nil" en todos
    sus bordes.

    Se usa "nil" (y no "none") porque "nil" tiene la semántica de
    "suprimir cualquier borde heredado del padre", incluyendo el tblBorders.
    Aun así, si la celda ADYACENTE tiene un borde explícito single, puede ganar.
    En la práctica, para cabeceras con bg_color aplicado esto produce el efecto
    visual correcto porque las celdas vacías del header se ven como extensión
    del color de fondo, sin borde visible molesto.
    """
    for row in tabla.rows:
        tcs = list(row._tr.findall(qn("w:tc")))
        for tc in tcs:
            texto = "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip()
            if not texto:
                # Usar "nil" para mayor potencia de supresión
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    tc.insert(0, tcPr)
                for old_b in tcPr.findall(qn("w:tcBorders")):
                    tcPr.remove(old_b)
                tcBorders = OxmlElement("w:tcBorders")
                for side in ("top", "left", "bottom", "right"):
                    el = OxmlElement(f"w:{side}")
                    el.set(qn("w:val"),   "nil")
                    el.set(qn("w:sz"),    "0")
                    el.set(qn("w:space"), "0")
                    el.set(qn("w:color"), "auto")
                    tcBorders.append(el)
                tcPr.append(tcBorders)
                _vaciar_celda_xml(tc)


def _crear_tabla_docx(doc, df, cell_formats, merges, es_loop=False):
    """
    Crea y retorna una tabla python-docx con datos, merges y formato.

    Pipeline:
    1. Crear tabla y aplicar bordes de cuadrícula.
    2. Volcar los datos en las celdas.
    3. Aplicar merges + formato (bg, bold, font_size, font_color).
    4. Limpiar párrafos vacíos extra generados por merge().
    5. Eliminar celdas vacías del final de cada fila ampliando la última
       celda con datos → tabla verdaderamente asimétrica sin bordes fantasma.
    6. Ocultar con val="nil" las celdas vacías restantes (no-trailing),
       típicamente cabeceras de tablas regulares con merges parciales.

    El parámetro es_loop se conserva por compatibilidad.
    """
    todas_las_filas = [list(df.columns)] + df.values.tolist()
    n_filas = len(todas_las_filas)
    n_cols  = len(todas_las_filas[0]) if todas_las_filas else 1

    tabla = doc.add_table(rows=n_filas, cols=n_cols)
    _aplicar_bordes_tabla(tabla)

    for fi, fila_datos in enumerate(todas_las_filas):
        for ci, val in enumerate(fila_datos):
            text = str(val) if val is not None and str(val) not in ("nan", "None") else ""
            tabla.cell(fi, ci).text = text

    _aplicar_formato_tabla(tabla, cell_formats, merges)

    # Paso 4: limpiar párrafos extra generados por merge()
    _limpiar_parrafos_merge(tabla)

    # Paso 5: eliminar celdas vacías trailing → tabla asimétrica real
    _hacer_filas_asimetricas(tabla)

    # Paso 6: solo en tablas de cantón — ocultar celdas vacías no-trailing.
    # Las tablas regulares conservan bordes uniformes en todas las celdas.
    if es_loop:
        _ocultar_celdas_vacias_intermedias(tabla)

    return tabla


# ==============================
# REEMPLAZO DE TABLAS
# ==============================

def reemplazar_tablas(doc_path, diccionario, _docs_service=None, drive_service=None):
    """
    Reemplaza placeholders de tipo table en el .docx local.
    Descarga cada Excel, construye la tabla python-docx (con merges y formato)
    e inserta en el lugar del placeholder.
    """
    doc          = Document(doc_path)
    reemplazados = []
    omitidos     = []

    open(os.path.join(os.path.dirname(os.path.abspath(__file__)), "debug_tablas.log"), "w").close()

    entradas = [(a, e) for a, e in diccionario.items()
                if e.get("type") == "table" and e.get("value") is not None]

    if not entradas:
        _log("No se encontraron entradas de tipo table.")
        return {"reemplazados": [], "omitidos": omitidos}

    for alias, entrada in entradas:
        placeholder = entrada.get("placeholder", "").strip()
        valor       = entrada.get("value", {})
        url_tabla   = valor.get("file_id")
        sheet_name  = valor.get("sheet")
        header_row  = valor.get("header_row")

        if not url_tabla:
            omitidos.append({"alias": alias, "razon": "URL de tabla vacia"})
            continue
        if not placeholder:
            omitidos.append({"alias": alias, "razon": "placeholder vacio"})
            continue

        try:
            df, cell_formats, merges = _cargar_datos_tabla(
                url_tabla, sheet_name, drive_service, header_row,
                alias=alias, placeholder=placeholder
            )
            if df is None or df.empty:
                omitidos.append({"alias": alias, "razon": "datos vacios"})
                continue

            n_filas    = len(df) + 1
            n_cols     = len(df.columns)
            ocurrencias = 0

            for para in list(_iter_all_paragraphs(doc)):
                _consolidar_runs(para)
                if placeholder not in _get_para_full_text(para):
                    continue
                ocurrencias += 1
                tabla = _crear_tabla_docx(doc, df, cell_formats, merges)
                para._element.addprevious(tabla._element)
                para._element.getparent().remove(para._element)
                _log(f"   '{alias}' ocurrencia {ocurrencias} -> {n_filas} filas x {n_cols} cols ({len(merges)} merges, {len(cell_formats)} formatos)")

            if ocurrencias > 0:
                reemplazados.append({"alias": alias, "n_filas": n_filas, "n_cols": n_cols, "ocurrencias": ocurrencias})
                if ocurrencias > 1:
                    _log(f"   '{alias}' -> {ocurrencias} ocurrencias reemplazadas en total")
            else:
                omitidos.append({"alias": alias, "razon": f"placeholder '{placeholder}' no encontrado"})

        except Exception as e:
            traceback.print_exc()
            omitidos.append({"alias": alias, "razon": f"error: {e}"})

    doc.save(doc_path)
    _log(f"Tablas reemplazadas: {len(reemplazados)}")
    if omitidos:
        _log(f"Tablas omitidas: {len(omitidos)}")
        for o in omitidos:
            _log(f"   {o['alias']}: {o['razon']}")
    return {"reemplazados": reemplazados, "omitidos": omitidos}


# ==============================
# LOOPS (tablas dinamicas por hojas)
# ==============================

def _listar_hojas_con_prefijo(url, prefijo, drive_service):
    """
    Descarga el Excel y retorna hojas cuyo nombre empiece con prefijo.
    Si prefijo es None/vacio, retorna todas.
    """
    file_id = extraer_id_gdoc(url)

    def _descargar(request):
        fh = io.BytesIO()
        dl = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            _, done = dl.next_chunk()
        fh.seek(0)
        return fh

    try:
        fh = _descargar(drive_service.files().get_media(fileId=file_id))
    except HttpError as e:
        if e.resp.status in (400, 403):
            fh = _descargar(drive_service.files().export_media(
                fileId=file_id,
                mimeType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ))
        else:
            raise

    wb = openpyxl.load_workbook(fh, read_only=True, data_only=True)
    hojas = list(wb.sheetnames)

    if not prefijo:
        return hojas
    prefijo_lower = prefijo.lower()
    return [h for h in hojas if h.lower().startswith(prefijo_lower)]


def reemplazar_loops(doc_path, diccionario, _docs_service=None, drive_service=None):
    """
    Reemplaza placeholders {% loop alias %} en el .docx local.
    Por cada hoja del Excel con el prefijo indicado inserta una tabla
    (con salto de pagina entre tablas, igual que la version original).
    Preserva merges y formato de celda.
    """
    doc          = Document(doc_path)
    reemplazados = []
    omitidos     = []

    entradas = [(a, e) for a, e in diccionario.items()
                if e.get("type") == "loop" and e.get("value") is not None]

    if not entradas:
        _log("No se encontraron entradas de tipo loop.")
        return {"reemplazados": [], "omitidos": omitidos}

    for alias, entrada in entradas:
        valor      = entrada["value"]
        url_excel  = valor.get("file_id")
        prefijo    = valor.get("prefijo")
        header_row = valor.get("header_row")
        placeholder = "{% loop " + alias + " %}"

        if not url_excel:
            omitidos.append({"alias": alias, "razon": "URL de Excel vacia"})
            continue

        try:
            hojas = _listar_hojas_con_prefijo(url_excel, prefijo, drive_service)

            if not hojas:
                omitidos.append({"alias": alias, "razon": f"No hay hojas con prefijo '{prefijo}'"})
                continue

            _log(f"   Loop '{alias}': {len(hojas)} hojas -> {hojas}")

            para_loop = None
            for para in list(_iter_all_paragraphs(doc)):
                _consolidar_runs(para)
                if placeholder in _get_para_full_text(para):
                    para_loop = para
                    break

            if para_loop is None:
                omitidos.append({"alias": alias, "razon": f"placeholder '{placeholder}' no encontrado"})
                continue

            ref_el            = para_loop._element
            tablas_insertadas = 0

            for hoja in hojas:
                try:
                    df, cell_formats, merges = _cargar_datos_tabla(
                        url_excel, hoja, drive_service, header_row,
                        alias=alias, placeholder=hoja
                    )
                    if df is None or df.empty:
                        _log(f"   Hoja '{hoja}' vacia, se omite.")
                        continue

                    n_filas = len(df) + 1
                    n_cols  = len(df.columns)

                    # Salto de pagina antes de cada tabla excepto la primera
                    if tablas_insertadas > 0:
                        para_salto = OxmlElement("w:p")
                        run_salto  = OxmlElement("w:r")
                        br         = OxmlElement("w:br")
                        br.set(qn("w:type"), "page")
                        run_salto.append(br)
                        para_salto.append(run_salto)
                        ref_el.addprevious(para_salto)

                    tabla = _crear_tabla_docx(doc, df, cell_formats, merges, es_loop=True)
                    ref_el.addprevious(tabla._element)

                    tablas_insertadas += 1
                    _log(f"   Tabla '{hoja}' insertada ({n_filas} x {n_cols})")

                except Exception as e_hoja:
                    traceback.print_exc()
                    _log(f"   Error en hoja '{hoja}': {e_hoja}")
                    omitidos.append({"alias": alias, "razon": f"error en hoja '{hoja}': {e_hoja}"})

            # Eliminar el parrafo placeholder
            para_loop._element.getparent().remove(para_loop._element)

            reemplazados.append({
                "alias": alias, "hojas": hojas, "tablas_insertadas": tablas_insertadas,
            })

        except Exception as e:
            traceback.print_exc()
            omitidos.append({"alias": alias, "razon": f"error: {e}"})

    doc.save(doc_path)
    _log(f"Loops reemplazados: {len(reemplazados)}")
    for r in reemplazados:
        _log(f"   {r['alias']} -> {r['tablas_insertadas']} tablas ({', '.join(r['hojas'])})")
    if omitidos:
        _log(f"Loops omitidos: {len(omitidos)}")
        for o in omitidos:
            _log(f"   {o['alias']}: {o['razon']}")
    return {"reemplazados": reemplazados, "omitidos": omitidos}


# ==============================
# UTILIDADES
# ==============================

def obtener_ruta_credenciales():
    ruta = os.getenv("GOOGLE_CREDS")
    if not ruta:
        raise Exception("Debes definir la variable de entorno GOOGLE_CREDS")
    if not os.path.exists(ruta):
        raise FileNotFoundError(f"No existe: {ruta}")
    return ruta


# ==============================
# TABLAS DESDE ARCHIVO LOCAL
# (para entornos sin Google Drive, ej. Colab con Drive montado)
# ==============================

def _cargar_datos_tabla_local(ruta_excel, sheet_name, header_row=None):
    """
    Lee una hoja de un .xlsx LOCAL y devuelve (df, cell_formats, merges).

    Equivalente a _cargar_datos_tabla() pero sin Google Drive:
    lee directamente desde una ruta en disco (o Drive montado en Colab).
    Usa openpyxl para respetar celdas fusionadas en los encabezados,
    evitando las columnas 'Unnamed: N' que genera pd.read_excel().

    Parametros
    ----------
    ruta_excel  : str        – ruta local al .xlsx
    sheet_name  : str|None   – nombre de la hoja; si es None usa la activa
    header_row  : int|None   – fila de encabezado (1-based). Si es None o 0
                               se detecta como la primera fila con contenido.
    """
    wb = openpyxl.load_workbook(ruta_excel, data_only=True)

    if sheet_name and sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
    else:
        ws = wb.active

    # Calcular rango real con datos
    max_col_real = max_row_real = 0
    first_row_con_datos = None
    for fila in ws.iter_rows():
        for celda in fila:
            val = celda.value
            if val is not None and str(val).strip() not in ("", "'"):
                max_col_real = max(max_col_real, celda.column)
                max_row_real = max(max_row_real, celda.row)
                if first_row_con_datos is None:
                    first_row_con_datos = celda.row
    if max_col_real == 0:
        max_col_real = ws.max_column or 1
    if max_row_real == 0:
        max_row_real = ws.max_row or 1

    # Determinar la fila de encabezado (0-based)
    if header_row and int(header_row) > 0:
        header_idx = int(header_row) - 1
        fila_c = list(ws.iter_rows(
            min_row=header_idx + 1, max_row=header_idx + 1, max_col=max_col_real
        ))
        if fila_c:
            tiene_cont = any(
                c.value is not None and str(c.value).strip() not in ("", "'")
                for c in fila_c[0]
            )
            if not tiene_cont and first_row_con_datos is not None:
                header_idx = first_row_con_datos - 1
    else:
        header_idx = (first_row_con_datos - 1) if first_row_con_datos else 0

    filas_excel = list(ws.iter_rows(
        min_row=header_idx + 1,
        max_row=max_row_real,
        max_col=max_col_real,
        values_only=False,
    ))
    if not filas_excel:
        return None, {}, []

    def _limpiar_enc(val):
        if val is None:
            return ""
        s = str(val)
        return s.lstrip("'") if s.startswith("'") else s

    encabezados = [_limpiar_enc(c.value) for c in filas_excel[0]]
    datos_filas = [
        [str(c.value) if c.value is not None else "" for c in fila]
        for fila in filas_excel[1:]
    ]
    df = pd.DataFrame(datos_filas, columns=encabezados).fillna("").astype(str)

    _FILLS_IGNORADOS = {"00000000", "FFFFFFFF", "FF000000", "00FFFFFF"}
    cell_formats = {}
    for doc_ri, fila in enumerate(filas_excel):
        for ci, celda in enumerate(fila):
            if ci >= max_col_real:
                break
            has_value = celda.value is not None and str(celda.value).strip() not in ("", "'")
            fmt = {}
            fill = celda.fill
            if fill and fill.fill_type not in (None, "none"):
                color = fill.fgColor
                if color and color.type == "rgb" and color.rgb not in _FILLS_IGNORADOS:
                    fmt["bg_color"] = color.rgb
            if has_value:
                if celda.font and celda.font.bold:
                    fmt["bold"] = True
                if celda.font and celda.font.size:
                    fmt["font_size"] = float(celda.font.size)
                if celda.font and celda.font.color:
                    fc = celda.font.color
                    if fc.type == "rgb" and fc.rgb not in ("00000000", "FF000000"):
                        fmt["font_color"] = fc.rgb
            if fmt:
                cell_formats[(doc_ri, ci)] = fmt

    merges = []
    for rng in ws.merged_cells.ranges:
        doc_min_row = rng.min_row - 1 - header_idx
        doc_max_row = rng.max_row - 1 - header_idx
        min_col     = rng.min_col - 1
        max_col     = min(rng.max_col - 1, max_col_real - 1)
        if doc_min_row >= 0 and min_col < max_col_real:
            merges.append({
                "min_row":  doc_min_row, "max_row":  doc_max_row,
                "min_col":  min_col,     "max_col":  max_col,
                "row_span": doc_max_row - doc_min_row + 1,
                "col_span": max_col - min_col + 1,
            })

    return df, cell_formats, merges


def reemplazar_tablas_local(doc_path, diccionario, ruta_excel_calculos):
    """
    Version de reemplazar_tablas() para archivos locales (sin Google Drive).
    Util en Colab cuando el Excel ya esta en disco o en Drive montado.

    Usa openpyxl para leer cada hoja, preservando celdas fusionadas
    y formato de celda (negrita, colores de fondo, tamaño de fuente).

    Parametros
    ----------
    doc_path             : str  – ruta local al .docx (se modifica in-place)
    diccionario          : dict – devuelto por cargar_diccionario()
    ruta_excel_calculos  : str  – ruta local al Excel con los datos de tabla

    Cada entrada del diccionario de tipo 'table' debe tener en 'value':
        sheet      : nombre de la hoja del Excel
        header_row : (opcional) fila de encabezado, 1-based
    """
    doc          = Document(doc_path)
    reemplazados = []
    omitidos     = []

    entradas = [
        (alias, e) for alias, e in diccionario.items()
        if e.get("type") == "table" and e.get("value") is not None
    ]

    if not entradas:
        _log("No se encontraron entradas de tipo table.")
        return {"reemplazados": [], "omitidos": omitidos}

    for alias, entrada in entradas:
        placeholder = entrada.get("placeholder", "").strip()
        valor       = entrada.get("value", {})
        sheet_name  = valor.get("sheet")
        header_row  = valor.get("header_row")

        if not placeholder:
            omitidos.append({"alias": alias, "razon": "placeholder vacio"})
            continue
        if not sheet_name:
            omitidos.append({"alias": alias, "placeholder": placeholder,
                             "razon": "sin sheet definido en el diccionario"})
            continue

        try:
            df, cell_formats, merges = _cargar_datos_tabla_local(
                ruta_excel_calculos, sheet_name, header_row
            )
        except Exception as e:
            traceback.print_exc()
            omitidos.append({"alias": alias, "placeholder": placeholder,
                             "razon": f"error leyendo sheet '{sheet_name}': {e}"})
            continue

        if df is None or df.empty:
            omitidos.append({"alias": alias, "placeholder": placeholder,
                             "razon": f"sheet '{sheet_name}' vacio"})
            continue

        ocurrencias = 0
        for para in list(_iter_all_paragraphs(doc)):
            _consolidar_runs(para)
            if placeholder not in _get_para_full_text(para):
                continue
            tabla = _crear_tabla_docx(doc, df, cell_formats, merges)
            para._element.addprevious(tabla._element)
            para._element.getparent().remove(para._element)
            ocurrencias += 1

        if ocurrencias > 0:
            reemplazados.append({
                "alias": alias, "sheet": sheet_name,
                "n_filas": len(df) + 1, "n_cols": len(df.columns),
                "merges": len(merges), "ocurrencias": ocurrencias,
            })
            _log(f"   '{alias}' -> sheet '{sheet_name}' "
                 f"({len(df)+1} filas x {len(df.columns)} cols, "
                 f"{len(merges)} merges, {len(cell_formats)} formatos)")
        else:
            omitidos.append({"alias": alias, "placeholder": placeholder,
                             "razon": f"placeholder '{placeholder}' no encontrado"})

    doc.save(doc_path)
    _log(f"Tablas locales reemplazadas: {len(reemplazados)}")
    if omitidos:
        _log(f"Tablas omitidas: {len(omitidos)}")
        for o in omitidos:
            _log(f"   {o.get('alias')}: {o.get('razon')}")
    return {"reemplazados": reemplazados, "omitidos": omitidos}


def reemplazar_loops_local(doc_path, diccionario, ruta_excel_calculos):
    """
    Version de reemplazar_loops() para archivos locales (sin Google Drive API).
    Reemplaza placeholders {% loop alias %} en el .docx.
    Por cada hoja del Excel cuyo nombre empiece con el prefijo definido
    en el diccionario inserta una tabla (con salto de pagina entre tablas).
    Preserva merges y formato de celda usando openpyxl.
    """
    doc          = Document(doc_path)
    reemplazados = []
    omitidos     = []

    entradas = [
        (alias, e) for alias, e in diccionario.items()
        if e.get("type") == "loop" and e.get("value") is not None
    ]

    if not entradas:
        _log("No se encontraron entradas de tipo loop.")
        return {"reemplazados": [], "omitidos": omitidos}

    wb_cache = {}  # cache para no abrir el mismo workbook varias veces

    for alias, entrada in entradas:
        valor       = entrada["value"]
        prefijo     = valor.get("prefijo")
        header_row  = valor.get("header_row")
        placeholder = "{% loop " + alias + " %}"

        try:
            # Listar hojas con el prefijo indicado desde el archivo local
            if ruta_excel_calculos not in wb_cache:
                wb_cache[ruta_excel_calculos] = openpyxl.load_workbook(
                    ruta_excel_calculos, read_only=True, data_only=True
                )
            wb    = wb_cache[ruta_excel_calculos]
            hojas = list(wb.sheetnames)
            if prefijo:
                prefijo_lower = prefijo.lower()
                hojas = [h for h in hojas if h.lower().startswith(prefijo_lower)]

            if not hojas:
                omitidos.append({"alias": alias,
                                 "razon": f"No hay hojas con prefijo '{prefijo}'"})
                continue

            _log(f"   Loop '{alias}': {len(hojas)} hojas -> {hojas}")

            para_loop = None
            for para in list(_iter_all_paragraphs(doc)):
                _consolidar_runs(para)
                if placeholder in _get_para_full_text(para):
                    para_loop = para
                    break

            if para_loop is None:
                omitidos.append({"alias": alias,
                                 "razon": f"placeholder '{placeholder}' no encontrado"})
                continue

            ref_el            = para_loop._element
            tablas_insertadas = 0

            for hoja in hojas:
                try:
                    df, cell_formats, merges = _cargar_datos_tabla_local(
                        ruta_excel_calculos, hoja, header_row
                    )
                    if df is None or df.empty:
                        _log(f"   Hoja '{hoja}' vacia, se omite.")
                        continue

                    if tablas_insertadas > 0:
                        para_salto = OxmlElement("w:p")
                        run_salto  = OxmlElement("w:r")
                        br         = OxmlElement("w:br")
                        br.set(qn("w:type"), "page")
                        run_salto.append(br)
                        para_salto.append(run_salto)
                        ref_el.addprevious(para_salto)

                    tabla = _crear_tabla_docx(doc, df, cell_formats, merges, es_loop=True)
                    ref_el.addprevious(tabla._element)
                    tablas_insertadas += 1
                    _log(f"   Tabla '{hoja}' insertada "
                         f"({len(df)+1} x {len(df.columns)}, {len(merges)} merges)")

                except Exception as e_hoja:
                    traceback.print_exc()
                    omitidos.append({"alias": alias,
                                     "razon": f"error en hoja '{hoja}': {e_hoja}"})

            para_loop._element.getparent().remove(para_loop._element)
            reemplazados.append({
                "alias": alias, "hojas": hojas,
                "tablas_insertadas": tablas_insertadas,
            })

        except Exception as e:
            traceback.print_exc()
            omitidos.append({"alias": alias, "razon": f"error: {e}"})

    doc.save(doc_path)
    _log(f"Loops locales reemplazados: {len(reemplazados)}")
    for r in reemplazados:
        _log(f"   {r['alias']} -> {r['tablas_insertadas']} tablas ({', '.join(r['hojas'])})")
    if omitidos:
        _log(f"Loops omitidos: {len(omitidos)}")
        for o in omitidos:
            _log(f"   {o.get('alias')}: {o.get('razon')}")
    return {"reemplazados": reemplazados, "omitidos": omitidos}

